"""
Document Processor - Extract text from uploaded files

WHY THIS EXISTS:
Documents come in different formats (PDF, DOCX, Markdown), but the rest
of our pipeline just needs plain text. This module is the "translator"
that normalizes everything into a common format.

HOW IT WORKS:
1. User uploads a file
2. We detect the file type from the extension
3. We use the right library to extract text:
   - PDF  → pdfplumber (reads each page, extracts text)
   - DOCX → python-docx (reads each paragraph)
   - MD   → just read the raw text (it's already text!)
4. Return the text + metadata (filename, page count, etc.)

WHAT WE RETURN:
A ProcessedDocument with:
  - text: the full extracted text
  - metadata: filename, file type, page count, character count
"""

import os
from dataclasses import dataclass, field
from pathlib import Path

import pdfplumber
from docx import Document as DocxDocument


@dataclass
class ProcessedDocument:
    """Container for extracted document content + metadata."""
    text: str
    metadata: dict = field(default_factory=dict)


def process_file(file_path: str) -> ProcessedDocument:
    """
    Main entry point. Takes a file path, returns extracted text.
    
    Args:
        file_path: Path to the uploaded file
        
    Returns:
        ProcessedDocument with text and metadata
        
    Raises:
        ValueError: If file type is not supported
    """
    path = Path(file_path)
    extension = path.suffix.lower()
    
    extractors = {
        ".pdf": _extract_pdf,
        ".docx": _extract_docx,
        ".md": _extract_markdown,
        ".txt": _extract_markdown,  # .txt uses same logic as .md
    }
    
    if extension not in extractors:
        supported = ", ".join(extractors.keys())
        raise ValueError(f"Unsupported file type: {extension}. Supported: {supported}")
    
    text, extra_metadata = extractors[extension](file_path)
    
    # Build metadata that every document gets
    metadata = {
        "filename": path.name,
        "file_type": extension,
        "char_count": len(text),
        "word_count": len(text.split()),
        **extra_metadata,
    }
    
    return ProcessedDocument(text=text, metadata=metadata)


def _extract_pdf(file_path: str) -> tuple[str, dict]:
    """
    Extract text from a PDF file.
    
    pdfplumber reads each page and extracts text, preserving layout
    better than most PDF libraries. We join pages with double newlines
    so the chunker can detect page boundaries later.
    """
    pages = []
    
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                pages.append(page_text.strip())
    
    text = "\n\n".join(pages)
    metadata = {"page_count": len(pages)}
    
    return text, metadata


def _extract_docx(file_path: str) -> tuple[str, dict]:
    """
    Extract text from a Word document.
    
    python-docx reads each paragraph element. We skip empty paragraphs
    and join with newlines. Headings and body text both come through
    as paragraphs — the chunker will handle structure later.
    """
    doc = DocxDocument(file_path)
    
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())
    
    text = "\n\n".join(paragraphs)
    metadata = {"paragraph_count": len(paragraphs)}
    
    return text, metadata


def _extract_markdown(file_path: str) -> tuple[str, dict]:
    """
    Extract text from Markdown or plain text files.
    
    These are already text, so we just read them. We do strip
    extra whitespace to keep things clean.
    """
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read().strip()
    
    metadata = {"line_count": len(text.splitlines())}
    
    return text, metadata


# ── Quick test ──────────────────────────────────────────────
# You can run this file directly to test it:
#   python -m app.document_processor
#
if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python -m app.document_processor <file_path>")
        print("Supported: .pdf, .docx, .md, .txt")
        sys.exit(1)
    
    result = process_file(sys.argv[1])
    print(f"\n📄 Processed: {result.metadata['filename']}")
    print(f"   Type: {result.metadata['file_type']}")
    print(f"   Words: {result.metadata['word_count']}")
    print(f"   Characters: {result.metadata['char_count']}")
    print(f"\n── First 500 chars ──")
    print(result.text[:500])
