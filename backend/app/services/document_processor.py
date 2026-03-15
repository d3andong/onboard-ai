"""
Document Processor - Extract text from uploaded files

Supports PDF (pdfplumber), DOCX (python-docx), Markdown, and plain text.
Returns a ProcessedDocument with full text and metadata.

Independently testable:
    python -m app.services.document_processor <file_path>
"""

from dataclasses import dataclass, field
from pathlib import Path

import pdfplumber
from docx import Document as DocxDocument


@dataclass
class ProcessedDocument:
    """Container for extracted document content and metadata."""
    text: str
    metadata: dict = field(default_factory=dict)


def process_file(file_path: str) -> ProcessedDocument:
    """
    Extract text from a file based on its extension.

    Args:
        file_path: Absolute path to the file.

    Returns:
        ProcessedDocument with .text and .metadata.

    Raises:
        ValueError: Unsupported file extension.
        IOError:    File cannot be read.
    """
    path = Path(file_path)
    extension = path.suffix.lower()

    extractors = {
        ".pdf":  _extract_pdf,
        ".docx": _extract_docx,
        ".md":   _extract_text,
        ".txt":  _extract_text,
    }

    if extension not in extractors:
        supported = ", ".join(extractors.keys())
        raise ValueError(f"Unsupported file type '{extension}'. Supported: {supported}")

    text, extra_meta = extractors[extension](file_path)

    metadata = {
        "filename":   path.name,
        "file_type":  extension,
        "char_count": len(text),
        "word_count": len(text.split()),
        **extra_meta,
    }

    return ProcessedDocument(text=text, metadata=metadata)


# ── Private extractors ────────────────────────────────────────

def _extract_pdf(file_path: str) -> tuple[str, dict]:
    pages = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                pages.append(page_text.strip())
    return "\n\n".join(pages), {"page_count": len(pages)}


def _extract_docx(file_path: str) -> tuple[str, dict]:
    doc = DocxDocument(file_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs), {"paragraph_count": len(paragraphs)}


def _extract_text(file_path: str) -> tuple[str, dict]:
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read().strip()
    return text, {"line_count": len(text.splitlines())}


# ── Quick test ───────────────────────────────────────────────
if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python -m app.services.document_processor <file_path>")
        sys.exit(1)

    result = process_file(sys.argv[1])
    print(f"File:    {result.metadata['filename']}")
    print(f"Type:    {result.metadata['file_type']}")
    print(f"Words:   {result.metadata['word_count']}")
    print(f"Chars:   {result.metadata['char_count']}")
    print(f"\n── First 500 chars ──\n{result.text[:500]}")
